from tqdm import tqdm
import pickle
from docx import Document
from typing import List, Dict, Any

pickleData = []

def extract_text_from_table(table) -> List[Dict[str, str]]:
    extracted_text = []
    for row in table.rows:
        for cell in row.cells[1:]:  # Skip the first column
            for para in cell.paragraphs:
                for run in para.runs:
                    text = run.text.strip()
                    if not text:
                        continue

                    if run.italic:
                        extracted_text.append({"speaker": "patient", "text": text})
                    else:
                        extracted_text.append({"speaker": "doctor", "text": text})
    return extracted_text

def extract_text_and_check_formatting(file_path: str) -> List[List[str]]:
    document = Document(file_path)
    extracted_text = []

    
    for table in document.tables:
        table_text = extract_text_from_table(table)
        extracted_text.extend(table_text)

    merged_text = []
    current_speaker = ""
    current_text = ""

    for entry in extracted_text:
        if entry["speaker"] == "doctor":
            if current_speaker == "patient":
                merged_text.append(["patient", current_text.strip()])
                current_text = ""
            current_speaker = "doctor"
            current_text += entry["text"] + " "
        elif entry["speaker"] == "patient":
            if current_speaker == "doctor":
                merged_text.append(["doctor", current_text.strip()])
                current_text = ""
            current_speaker = "patient"
            current_text += entry["text"] + " "

    if current_text:
        merged_text.append([current_speaker, current_text.strip()])

    return merged_text

file_path = '/data/MBESLIS/data/transfer_2713523_files_60e044cf/Transcripten ABEL/1001_A.docx'
pickleData.append(extract_text_and_check_formatting(file_path))


file_path = '/data/MBESLIS/data/transfer_2713523_files_60e044cf/Transcripten ABEL/1002_A.docx'
pickleData.append(extract_text_and_check_formatting(file_path))


file_path = '/data/MBESLIS/data/transfer_2713523_files_60e044cf/Transcripten ABEL/1003_A.docx'
pickleData.append(extract_text_and_check_formatting(file_path))

file_path = '/data/MBESLIS/data/transfer_2713523_files_60e044cf/Transcripten ABEL/1004 AB_A.docx'
pickleData.append(extract_text_and_check_formatting(file_path))

file_path = '/data/MBESLIS/data/transfer_2713523_files_60e044cf/Transcripten ABEL/1035_A.docx'
pickleData.append(extract_text_and_check_formatting(file_path))


with open('/data/MBESLIS/MBESLIS.pkl', 'wb') as file:
    pickle.dump(pickleData, file)
    
