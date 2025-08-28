import pickle
#with open('/data/MBESLIS/MBESLIS.pkl', 'rb') as file:
   #pickleData = pickle.load(file)

import os
from docx import Document



def parse_conversation(doc_path):
    doc = Document(doc_path)
    conversations = []
    current_speaker = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        if not text or text.startswith("Consult:"):
            continue
        
        if any(run.bold for run in para.runs):
            current_speaker = 'doctor'
        elif any(run.italic for run in para.runs):
            current_speaker = 'patient'
        else:
            if current_speaker is None:
                continue
        
        if text.startswith("V: "):
            text = text.replace("V: ", "")

        conversations.append([current_speaker, text])
    
    return conversations



errorcounter = 0
working = 0
for file_name in os.listdir("/home/s3366952/master-thesis/datasets/MBESLIS/data/transfer_2713631_files_007eaaac/"):
    absolute_path = os.path.join("/home/s3366952/master-thesis/datasets/MBESLIS/data/transfer_2713631_files_007eaaac/", file_name)
    if os.path.isfile(absolute_path):
        try:
            doc = Document(absolute_path)
            working += 1
        except:
            errorcounter += 1



for file_name in os.listdir("/home/s3366952/master-thesis/datasets/MBESLIS/data/transfer_2713661_files_90256190/"):
    absolute_path = os.path.join("/home/s3366952/master-thesis/datasets/MBESLIS/data/transfer_2713661_files_90256190/", file_name)
    if os.path.isfile(absolute_path):
        try:
            doc = Document(absolute_path)
            working += 1
        except:
            errorcounter += 1


for file_name in os.listdir("/home/s3366952/master-thesis/datasets/MBESLIS/data/transfer_2713680_files_27a4456a/"):
    absolute_path = os.path.join("/home/s3366952/master-thesis/datasets/MBESLIS/data/transfer_2713680_files_27a4456a/", file_name)
    if os.path.isfile(absolute_path):
        try:
            doc = Document(absolute_path)
            working += 1
        except:
            errorcounter += 1


print("Error", errorcounter, "Working", working)









#with open('/data/MBESLIS/MBESLIS.pkl', 'wb') as file:
    #pickle.dump(pickleData, file)